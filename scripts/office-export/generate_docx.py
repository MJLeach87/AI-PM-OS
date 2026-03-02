"""
DOCX generator for PM OS Office Export.
Converts a Markdown file to a branded Word document using python-docx.

Branding spec: pm-os-reference/identity/BRANDING.md
Constants:      scripts/office-export/branding.py
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import branding as B


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _rgb(hex_color: str) -> RGBColor:
    r, g, b = B.hex_to_rgb(hex_color)
    return RGBColor(r, g, b)


def _set_cell_bg(cell, hex_color: str):
    """Apply a solid background fill to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _apply_heading_style(para, level: int):
    """Apply PM OS heading style to a paragraph."""
    size_map = {1: B.FONT_SIZE_H1, 2: B.FONT_SIZE_H2, 3: B.FONT_SIZE_H3}
    bold_map = {1: True, 2: True, 3: True}

    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.font.name = B.FONT_PRIMARY
    run.font.size = Pt(size_map.get(level, B.FONT_SIZE_H3))
    run.font.bold = bold_map.get(level, False)
    run.font.color.rgb = _rgb(B.TEXT_PRIMARY)

    para.paragraph_format.space_after = Pt(B.HEADING_SPACE_AFTER_PT)


def _apply_body_style(para):
    for run in para.runs:
        run.font.name = B.FONT_PRIMARY
        run.font.size = Pt(B.FONT_SIZE_BODY)
        run.font.color.rgb = _rgb(B.TEXT_SECONDARY)


def _apply_code_style(para):
    """Monospace code block paragraph."""
    for run in para.runs:
        run.font.name = B.FONT_MONO
        run.font.size = Pt(B.FONT_SIZE_MONO)
        run.font.color.rgb = _rgb(B.TEXT_PRIMARY)
    # Light background shading on the paragraph
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), B.SURFACE_ALT.lstrip('#'))
    pPr.append(shd)


# ---------------------------------------------------------------------------
# Markdown parser (line-level, sufficient for PM OS artifact structure)
# ---------------------------------------------------------------------------

def _parse_markdown(md_text: str) -> list[dict]:
    """
    Tokenise Markdown into a simple token list.
    Token keys: type, level (headings), text, rows (tables), header (tables)
    """
    tokens = []
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith('```'):
            if in_code:
                tokens.append({'type': 'code', 'text': '\n'.join(code_lines)})
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            tokens.append({'type': 'heading', 'level': len(m.group(1)), 'text': m.group(2)})
            i += 1
            continue

        # Table — collect all consecutive table lines
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # First row = header, second row = separator (skip), rest = data
            header = [c.strip() for c in table_lines[0].split('|') if c.strip()]
            rows = []
            for tl in table_lines[2:]:
                row = [c.strip() for c in tl.split('|') if c.strip()]
                if row:
                    rows.append(row)
            tokens.append({'type': 'table', 'header': header, 'rows': rows})
            continue

        # Blank line
        if line.strip() == '':
            i += 1
            continue

        # Bullet / numbered list items — treat as body paragraph
        body_text = re.sub(r'^[\-\*\+]\s+', '', line)
        body_text = re.sub(r'^\d+\.\s+', '', body_text)
        # Strip inline markdown (bold, italic, code, links)
        body_text = re.sub(r'\*\*(.+?)\*\*', r'\1', body_text)
        body_text = re.sub(r'\*(.+?)\*', r'\1', body_text)
        body_text = re.sub(r'`(.+?)`', r'\1', body_text)
        body_text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', body_text)
        tokens.append({'type': 'paragraph', 'text': body_text})
        i += 1

    return tokens


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def _set_page_margins(doc: Document):
    for section in doc.sections:
        margin = Inches(B.PAGE_MARGIN_IN)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def _add_table(doc: Document, header: list[str], rows: list[list[str]]):
    col_count = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for j, cell in enumerate(hdr_cells):
        cell.text = header[j] if j < len(header) else ''
        _set_cell_bg(cell, B.PRIMARY)
        for run in cell.paragraphs[0].runs:
            run.font.name = B.FONT_PRIMARY
            run.font.size = Pt(B.FONT_SIZE_BODY)
            run.font.bold = True
            run.font.color.rgb = _rgb(B.SURFACE)

    # Data rows with alternating background
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        bg = B.SURFACE if i % 2 == 0 else B.SURFACE_ALT
        for j, cell in enumerate(row_cells):
            cell.text = row_data[j] if j < len(row_data) else ''
            _set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.name = B.FONT_PRIMARY
                run.font.size = Pt(B.FONT_SIZE_BODY)
                run.font.color.rgb = _rgb(B.TEXT_SECONDARY)

    doc.add_paragraph()  # space after table


def generate(source_path: Path, out_dir: Path) -> Path:
    """Convert a Markdown file to a branded DOCX. Returns the output path."""
    md_text = source_path.read_text(encoding='utf-8')
    tokens = _parse_markdown(md_text)

    doc = Document()
    _set_page_margins(doc)

    for token in tokens:
        t = token['type']

        if t == 'heading':
            level = token['level']
            style_name = f'Heading {min(level, 3)}'
            try:
                para = doc.add_heading(token['text'], level=min(level, 3))
            except Exception:
                para = doc.add_paragraph(token['text'])
            _apply_heading_style(para, level)

        elif t == 'paragraph':
            if not token['text'].strip():
                continue
            para = doc.add_paragraph(token['text'])
            _apply_body_style(para)

        elif t == 'code':
            # Bordered placeholder for code / Mermaid blocks
            para = doc.add_paragraph(token['text'])
            _apply_code_style(para)

        elif t == 'table':
            _add_table(doc, token['header'], token['rows'])

    stem = source_path.stem
    out_path = out_dir / f'{stem}.docx'
    doc.save(str(out_path))
    return out_path
